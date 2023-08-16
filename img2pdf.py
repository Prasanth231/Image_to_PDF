import os
import telegram
from telegram.ext import Updater, MessageHandler, Filters
import cv2
import pytesseract
import warnings
from docx import Document
from docx2pdf import convert
import time
from time import sleep

def receive_image(update, context):
    pythoncom.CoInitialize()
    
    chat_id = update.effective_chat.id
    message_id = update.message.message_id
    #print("level 1")

    if update.message.photo:

        photo_file_id = update.message.photo[-1].file_id

        photo_file = context.bot.get_file(photo_file_id)
        filename = f"{message_id}.jpeg"
        photo_file.download(filename)
        print("downloaded")
        context.bot.send_message(chat_id=chat_id, text="Thanks for the image!")
    else:
        def error(update, context):
            logger.error('Update "%s" caused error "%s"', update, context.error)
            context.bot.send_message(chat_id=chat_id, text='An error occurred. Please try again later.')

    pytesseract.pytesseract.tesseract_cmd=r'C:\\Users\\family\\AppData\\Local\\Programs\\Tesseract-OCR\\tesseract.exe'
    img= cv2.imread(filename)
    imgtext= pytesseract.image_to_string(filename)
    context.bot.send_message(chat_id=chat_id, text=imgtext)
    doc_file= 'Text.docx'
    document = Document()
    document.add_paragraph(imgtext)
    document.save(doc_file)
    with open(doc_file, 'rb') as f:
        context.bot.send_document(chat_id=chat_id, document=f, filename='Text.docx')
    pdf_file= "Text.pdf"
    #pdf_file = doc_path.replace(".docx", ".pdf")
    #convert(doc_file)
    convert(doc_file, pdf_file)
    with open(pdf_file, 'rb') as f:
        context.bot.send_document(chat_id=chat_id, document=f,filename='Text.pdf')
    time.sleep(5)
    if os.path.isfile(doc_file):
        os.remove(doc_file)
    if os.path.isfile(pdf_file):
        os.remove(pdf_file)
    if os.path.isfile(filename):
        os.remove(filename)        

bot_token = "6211989392:AAF04yaB3IP0fKeacWm-mPnLrLqORRWg9fQ"
updater = Updater(bot_token, use_context=True)
dispatcher = updater.dispatcher
image_handler = MessageHandler(Filters.photo, receive_image)
dispatcher.add_handler(image_handler)
updater.start_polling()
updater.idle()

'''import telegram
from telegram.ext import Updater, MessageHandler, Filters
import cv2
import warnings
from docx import Document
from docx2pdf import convert
import pythoncom
import time
from time import sleep
import requests
import os
import pytesseract


# Define a function to handle incoming image messages
def receive_image(update, context):
    pythoncom.CoInitialize()
    # Get the chat ID and message ID
    chat_id = update.effective_chat.id
    message_id = update.message.message_id
    #print("level 1")
    # Check if the message contains a photo
    if update.message.photo:
        # Get the photo file ID
        photo_file_id = update.message.photo[-1].file_id

        # Download the photo and save it to disk
        photo_file = context.bot.get_file(photo_file_id)
        filename = f"{message_id}.jpeg"
        photo_file.download(filename)
        print("downloaded")

        # Reply to the user with a confirmation message
        context.bot.send_message(chat_id=chat_id, text="Thanks for the image!")
    else:
        def error(update, context):
            logger.error('Update "%s" caused error "%s"', update, context.error)
            context.bot.send_message(chat_id=chat_id, text='An error occurred. Please try again later.')
    pytesseract.pytesseract.tesseract_cmd = os.path.join(os.getcwd(), "tesseract", "tesseract.exe")

    #url = 'https://github.com/UB-Mannheim/tesseract/wiki/tesseract-ocr-w64-setup-5.3.0.20221222.exe'
    #response = requests.get(url)

    # Save the file to disk
    #ith open('tesseract.exe', 'wb') as f:
        #f.write(response.content)

    #    Set the path to the executable
    #pytesseract.pytesseract.tesseract_cmd = os.path.abspath('tesseract.exe')
    #pytesseract.pytesseract.tesseract_cmd=r'https://drive.google.com/file/d/1O0ILsLXlWHCGQGzRve3OwcfaaDgClsOT/view?usp=share_link'
    img= cv2.imread(filename)
    imgtext= pytesseract.image_to_string(filename)
    context.bot.send_message(chat_id=chat_id, text=imgtext)
    doc_file= 'Text.docx'
    document = Document()
    document.add_paragraph(imgtext)
    document.save(doc_file)
    with open(doc_file, 'rb') as f:
        context.bot.send_document(chat_id=chat_id, document=f, filename='Text.docx')
    pdf_file= "Text.pdf"
    #pdf_file = doc_path.replace(".docx", ".pdf")
    #convert(doc_file)
    convert(doc_file, pdf_file)
    with open(pdf_file, 'rb') as f:
        context.bot.send_document(chat_id=chat_id, document=f,filename='Text.pdf')
    time.sleep(5)
    if os.path.isfile(doc_file):
        os.remove(doc_file)
    if os.path.isfile(pdf_file):
        os.remove(pdf_file)
    if os.path.isfile(filename):
        os.remove(filename)        

# Set up the Telegram bot
bot_token = "6211989392:AAF04yaB3IP0fKeacWm-mPnLrLqORRWg9fQ"
updater = Updater(bot_token, use_context=True)
dispatcher = updater.dispatcher

# Set up a message handler to receive image messages
image_handler = MessageHandler(Filters.photo, receive_image)
dispatcher.add_handler(image_handler)


# Start the bot
updater.start_polling()

# Keep the bot running until interrupted
updater.idle()
'''





#bot = telegram.Bot(token='6211989392:AAF04yaB3IP0fKeacWm-mPnLrLqORRWg9fQ')

'''import cv2
import pytesseract

pytesseract.pytesseract.tesseract_cmd=r'C:\\Users\\family\\AppData\\Local\\Programs\\Tesseract-OCR\\tesseract.exe'
img= cv2.imread('img.jpeg')
text= pytesseract.image_to_string(img)
print(text)
'''